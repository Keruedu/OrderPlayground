# -*- coding: utf-8 -*-
import json
from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\playgrounds\order-playground")
OUT = ROOT / "docs" / "note2node"
ASSETS = OUT / "assets"


def font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def wrap(draw, text, width, fnt):
    words = text.split()
    lines, line = [], ""
    for word in words:
        candidate = (line + " " + word).strip()
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_box(draw, xy, text, fill, outline="#24404f", text_fill="#10212a", title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    fnt = font(24 if title else 20, bold=title)
    lines = wrap(draw, text, x2 - x1 - 26, fnt)
    total_h = len(lines) * (fnt.size + 5)
    y = y1 + ((y2 - y1 - total_h) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill=text_fill, font=fnt)
        y += fnt.size + 5


def arrow(draw, start, end, color="#25495c"):
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    # simple arrow head
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - sign * 14, y2 - 8), (x2 - sign * 14, y2 + 8)]
    else:
        sign = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - sign * 14), (x2 + 8, y2 - sign * 14)]
    draw.polygon(pts, fill=color)


def make_architecture(path):
    img = Image.new("RGB", (1600, 950), "#f6f8f7")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "Order Playground - High-level Architecture", fill="#10212a", font=font(36, True))
    boxes = {
        "User / Postman": (80, 180, 330, 300),
        "Keycloak\nLogin + Token": (80, 430, 330, 560),
        "Gateway API\nGo Gin": (520, 300, 830, 440),
        "MySQL\norders, items,\nworkflow_runs": (1030, 160, 1340, 300),
        "MongoDB\naudit_events,\nrequest_logs": (1030, 360, 1340, 500),
        "Temporal\nCreateOrderWorkflow": (520, 610, 850, 750),
        "Inventory gRPC": (1030, 610, 1340, 720),
        "Notifier gRPC": (1030, 770, 1340, 880),
    }
    fills = ["#d9ecff", "#fff0c7", "#dff4e6", "#efe5ff", "#ffe3df", "#dff0ff", "#eef5d3", "#fce5f1"]
    for (text, xy), fill in zip(boxes.items(), fills):
        draw_box(d, xy, text, fill, title=("Gateway" in text or "Temporal" in text))
    arrow(d, (330, 240), (520, 360))
    arrow(d, (200, 430), (200, 300))
    arrow(d, (830, 340), (1030, 230))
    arrow(d, (830, 400), (1030, 430))
    arrow(d, (680, 440), (680, 610))
    arrow(d, (850, 660), (1030, 665))
    arrow(d, (850, 710), (1030, 825))
    arrow(d, (1030, 250), (850, 640))
    arrow(d, (1030, 450), (850, 700))
    d.text((60, 900), "Public API: HTTP/JSON | Internal calls: gRPC over HTTP/2 | Workflow orchestration: Temporal", fill="#315064", font=font(22))
    img.save(path)


def make_sequence(path):
    img = Image.new("RGB", (1600, 980), "#fbfbf8")
    d = ImageDraw.Draw(img)
    d.text((60, 38), "Create Order - Build -> Debug -> Learn Flow", fill="#10212a", font=font(36, True))
    lanes = ["User", "Keycloak", "Gateway", "MySQL", "MongoDB", "Temporal", "gRPC Services"]
    x_positions = [110, 340, 570, 800, 1030, 1260, 1480]
    for x, lane in zip(x_positions, lanes):
        d.text((x - 60, 110), lane, fill="#183140", font=font(20, True))
        d.line([(x, 150), (x, 910)], fill="#ccd6db", width=2)
    steps = [
        (0, 1, 190, "Login"),
        (1, 0, 250, "Access token"),
        (0, 2, 320, "POST /api/orders"),
        (2, 2, 380, "Verify JWT + role"),
        (2, 3, 450, "Insert PENDING"),
        (2, 4, 520, "ORDER_CREATED"),
        (2, 5, 590, "Start workflow"),
        (5, 6, 660, "Reserve / Notify"),
        (5, 3, 730, "Update COMPLETED"),
        (5, 4, 800, "Audit final event"),
    ]
    for a, b, y, label in steps:
        if a == b:
            d.arc((x_positions[a] - 35, y - 20, x_positions[a] + 35, y + 30), 270, 90, fill="#25495c", width=3)
            d.text((x_positions[a] + 45, y - 8), label, fill="#10212a", font=font(18))
        else:
            arrow(d, (x_positions[a], y), (x_positions[b], y))
            mid = (x_positions[a] + x_positions[b]) // 2
            d.rectangle((mid - 88, y - 25, mid + 88, y - 3), fill="#fbfbf8")
            d.text((mid - 82, y - 27), label, fill="#10212a", font=font(17))
    img.save(path)


def make_keycloak(path):
    img = Image.new("RGB", (1500, 840), "#f7faf9")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Keycloak Mental Model", fill="#10212a", font=font(36, True))
    draw_box(d, (90, 170, 620, 360), "master realm\nRealm quản trị Keycloak\nCó internal client: order-playground-realm", "#fff0c7", title=True)
    draw_box(d, (850, 170, 1380, 360), "order-playground realm\nRealm của ứng dụng\nCó client business: gateway-api", "#d9ecff", title=True)
    draw_box(d, (850, 480, 1120, 650), "Users\nuser1\nadmin1", "#dff4e6")
    draw_box(d, (1160, 480, 1380, 650), "Roles\nuser\nadmin", "#efe5ff")
    arrow(d, (620, 265), (850, 265))
    d.text((660, 232), "admin có thể quản lý", fill="#315064", font=font(19))
    d.text((90, 735), "Mình không dùng order-playground-realm để tích hợp app. Gateway dùng client gateway-api trong realm order-playground.", fill="#315064", font=font(22))
    img.save(path)


def make_temporal(path):
    img = Image.new("RGB", (1500, 920), "#fbfbf8")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Temporal Workflow - Happy Path and Failure Path", fill="#10212a", font=font(36, True))
    y = 180
    steps = [
        ("ORDER_CREATED", "#d9ecff"),
        ("ValidateOrderActivity", "#dff4e6"),
        ("ReserveInventoryActivity", "#dff4e6"),
        ("ApprovePaymentActivity", "#dff4e6"),
        ("SendNotificationActivity", "#dff4e6"),
        ("ORDER_COMPLETED", "#c9efd4"),
    ]
    xs = [80, 310, 580, 850, 1120, 1280]
    widths = [190, 220, 230, 230, 230, 190]
    for i, ((label, fill), x, w) in enumerate(zip(steps, xs, widths)):
        draw_box(d, (x, y, x + w, y + 110), label, fill)
        if i < len(steps) - 1:
            arrow(d, (x + w, y + 55), (xs[i + 1], y + 55))
    draw_box(d, (580, 470, 900, 600), "Failure case\ninventory reject\nquantity > 5", "#ffe3df", title=True)
    draw_box(d, (1010, 470, 1300, 600), "ORDER_FAILED\nwrite audit event", "#ffd5cf", title=True)
    arrow(d, (695, 290), (695, 470))
    arrow(d, (900, 535), (1010, 535))
    d.text((80, 760), "Bài học debug: nếu activity name không khớp tên worker đăng ký, workflow sẽ retry mãi và UI cho thấy Pending Activity.", fill="#315064", font=font(23))
    img.save(path)


def make_debug_tree(path):
    img = Image.new("RGB", (1500, 950), "#f8fbfa")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Debug Decision Tree", fill="#10212a", font=font(36, True))
    draw_box(d, (560, 140, 940, 250), "Có lỗi khi demo?", "#d9ecff", title=True)
    leaves = [
        ((80, 360, 390, 500), "401 / 403\nMở Keycloak\nCheck token, role, realm", "#fff0c7"),
        ((430, 360, 740, 500), "Order chưa xong\nMở Temporal UI\nCheck workflow/activity", "#dff4e6"),
        ((780, 360, 1090, 500), "Data sai/mất\nMở MySQL/Mongo\nCheck order/events", "#efe5ff"),
        ((1130, 360, 1440, 500), "App không lên\nMở Docker/logs\nCheck ports, health", "#ffe3df"),
    ]
    for xy, label, fill in leaves:
        draw_box(d, xy, label, fill)
        arrow(d, (750, 250), ((xy[0] + xy[2]) // 2, xy[1]))
    draw_box(d, (250, 660, 610, 790), "Issuer mismatch\nToken issuer: localhost\nJWKS: keycloak hostname", "#f7e1ff")
    draw_box(d, (890, 660, 1250, 790), "Activity mismatch\nWorkflow gọi sai tên\nWorker không nhận task", "#ffdede")
    arrow(d, (235, 500), (430, 660))
    arrow(d, (585, 500), (1070, 660))
    img.save(path)


def make_terminal_panel(path, title, lines):
    img = Image.new("RGB", (1500, 900), "#111827")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((30, 30, 1470, 870), radius=20, fill="#0b1220", outline="#314155", width=2)
    d.ellipse((65, 65, 85, 85), fill="#ff5f56")
    d.ellipse((95, 65, 115, 85), fill="#ffbd2e")
    d.ellipse((125, 65, 145, 85), fill="#27c93f")
    d.text((180, 55), title, fill="#e5edf5", font=font(26, True))
    y = 130
    for line in lines:
        d.text((70, y), line, fill="#d1e7dd" if line.startswith("OK") else "#e5edf5", font=font(22))
        y += 38
    img.save(path)


def make_dashboard_panel(path, title, columns, rows, accent="#d9ecff"):
    img = Image.new("RGB", (1500, 900), "#f4f7f7")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((35, 35, 1465, 865), radius=18, fill="#ffffff", outline="#cbd7dd", width=2)
    d.rectangle((35, 35, 1465, 115), fill=accent)
    d.text((70, 60), title, fill="#10212a", font=font(30, True))
    y = 170
    x0 = 80
    widths = [360, 360, 580][: len(columns)]
    x = x0
    for col, w in zip(columns, widths):
        d.text((x, y), col, fill="#10212a", font=font(22, True))
        x += w
    d.line((70, y + 38, 1430, y + 38), fill="#ccd6db", width=2)
    y += 65
    for row in rows:
        x = x0
        d.rounded_rectangle((65, y - 12, 1435, y + 55), radius=8, fill="#f9fbfb", outline="#e3eaee")
        for value, w in zip(row, widths):
            d.text((x, y), value, fill="#203845", font=font(20))
            x += w
        y += 82
    img.save(path)


def make_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    mermaids = {
        "architecture.mmd": dedent("""
            flowchart TD
                User[User / Postman] --> Keycloak[Keycloak: Login + Token]
                Keycloak --> User
                User --> Gateway[Go Gin Gateway API]
                Gateway --> MySQL[(MySQL: orders, items, workflow_runs)]
                Gateway --> Mongo[(MongoDB: audit_events, request_logs)]
                Gateway --> Temporal[Temporal: CreateOrderWorkflow]
                Temporal --> Inventory[Inventory Service: gRPC]
                Temporal --> Notifier[Notifier Service: gRPC]
                Temporal --> MySQL
                Temporal --> Mongo
        """).strip(),
        "create-order-sequence.mmd": dedent("""
            sequenceDiagram
                participant U as User
                participant K as Keycloak
                participant G as Gateway API
                participant M as MySQL
                participant D as MongoDB
                participant T as Temporal
                participant I as Inventory gRPC
                participant N as Notifier gRPC
                U->>K: Login lấy access token
                K-->>U: Access token
                U->>G: POST /api/orders + Bearer token
                G->>G: Verify JWT và role
                G->>M: Insert order PENDING
                G->>D: Insert ORDER_CREATED
                G->>T: Start CreateOrderWorkflow
                T->>I: Reserve inventory
                T->>N: Send notification
                T->>M: Update order COMPLETED / FAILED
                T->>D: Insert final audit event
        """).strip(),
        "debug-tree.mmd": dedent("""
            flowchart TD
                A[Có lỗi khi demo?]
                A --> B[401/403: mở Keycloak]
                A --> C[Order chưa xong: mở Temporal UI]
                A --> D[Data sai: mở MySQL/MongoDB]
                A --> E[App không lên: mở Docker logs]
                B --> F[Check realm, client, token, role]
                C --> G[Check workflow history, activity retry]
                D --> H[Check orders, workflow_runs, audit_events]
                E --> I[Check port conflict, healthcheck, startup retry]
        """).strip(),
        "keycloak-model.mmd": dedent("""
            flowchart LR
                Master[master realm<br/>quản trị Keycloak]
                Internal[order-playground-realm<br/>internal admin client]
                AppRealm[order-playground realm<br/>realm của app]
                Gateway[gateway-api<br/>business client]
                Users[user1/admin1]
                Master --> Internal
                Master --> AppRealm
                AppRealm --> Gateway
                AppRealm --> Users
        """).strip(),
        "temporal-retry.mmd": dedent("""
            flowchart TD
                A[CreateOrderWorkflow] --> B[ValidateOrderActivity]
                B --> C[ReserveInventoryActivity]
                C -->|success| D[ApprovePaymentActivity]
                C -->|failure after retry| F[MarkOrderFailedActivity]
                D --> E[SendNotificationActivity]
                E --> G[MarkOrderCompletedActivity]
                F --> H[ORDER_FAILED audit]
                G --> I[ORDER_COMPLETED audit]
        """).strip(),
    }
    for name, content in mermaids.items():
        (ASSETS / name).write_text(content, encoding="utf-8")
    make_architecture(ASSETS / "architecture.png")
    make_sequence(ASSETS / "create-order-sequence.png")
    make_keycloak(ASSETS / "keycloak-model.png")
    make_temporal(ASSETS / "temporal-flow.png")
    make_debug_tree(ASSETS / "debug-tree.png")
    make_terminal_panel(
        ASSETS / "smoke-test-summary.png",
        "Smoke test summary",
        [
            "OK GET /healthz -> { status: ok }",
            "OK user1 token from Keycloak",
            "OK POST /api/orders -> status PENDING + workflow_id",
            "OK Temporal workflow -> COMPLETED",
            "OK audit events -> ORDER_CREATED, ORDER_VALIDATED, INVENTORY_RESERVED, PAYMENT_APPROVED, ORDER_COMPLETED",
            "",
            "Ports: 8080 Gateway | 8081 Keycloak | 8088 Temporal UI | 3306 MySQL | 27017 MongoDB",
        ],
    )
    make_terminal_panel(
        ASSETS / "debug-incidents.png",
        "Debug incidents captured from playground",
        [
            "PowerShell: running scripts is disabled -> use ExecutionPolicy Bypass or inline command",
            "Docker daemon not running -> docker API pipe not found",
            "MySQL GUI/JDBC -> Public Key Retrieval is not allowed",
            "Keycloak -> token issuer localhost != Docker hostname",
            "Temporal -> namespace default not ready during startup",
            "Temporal -> activity name mismatch caused retry loop",
        ],
    )
    make_dashboard_panel(
        ASSETS / "docker-containers.png",
        "Docker containers used by the playground",
        ["Container", "Port", "Role"],
        [
            ("order-playground-gateway", "8080", "Go Gin public API"),
            ("order-playground-keycloak", "8081", "Auth UI + token endpoint"),
            ("order-playground-temporal-ui", "8088", "Workflow observation UI"),
            ("order-playground-mysql", "3306", "Business data"),
            ("order-playground-mongodb", "27017", "Audit and request logs"),
            ("inventory/notifier", "9091/9092", "Internal gRPC services"),
        ],
        "#dff4e6",
    )
    make_dashboard_panel(
        ASSETS / "keycloak-realm-screen.png",
        "Keycloak realm view",
        ["Area", "Value", "Meaning"],
        [
            ("Realm", "order-playground", "Realm dùng cho app demo"),
            ("Client", "gateway-api", "Business client để lấy/verify token"),
            ("Users", "user1, admin1", "Tài khoản test API"),
            ("Roles", "user, admin", "Phân quyền route thường/admin"),
        ],
        "#fff0c7",
    )
    make_dashboard_panel(
        ASSETS / "temporal-completed-screen.png",
        "Temporal workflow observation",
        ["Workflow", "Status", "Events"],
        [
            ("CreateOrderWorkflow", "COMPLETED", "ORDER_CREATED"),
            ("ValidateOrderActivity", "COMPLETED", "ORDER_VALIDATED"),
            ("ReserveInventoryActivity", "COMPLETED", "INVENTORY_RESERVED"),
            ("ApprovePaymentActivity", "COMPLETED", "PAYMENT_APPROVED"),
            ("SendNotificationActivity", "COMPLETED", "ORDER_COMPLETED"),
        ],
        "#d9ecff",
    )
    make_dashboard_panel(
        ASSETS / "keycloak-clients-screen.png",
        "Keycloak clients that caused confusion",
        ["Realm", "Client", "Meaning"],
        [
            ("master", "order-playground-realm", "Internal admin client do Keycloak tạo"),
            ("order-playground", "gateway-api", "Client business thật của playground"),
            ("master", "security-admin-console", "Admin console client"),
            ("order-playground", "account-console", "Account management UI"),
        ],
        "#efe5ff",
    )
    make_dashboard_panel(
        ASSETS / "db-audit-events-screen.png",
        "Audit events query result",
        ["Order event", "Source", "What it proves"],
        [
            ("ORDER_CREATED", "gateway-api", "API nhận request và ghi MongoDB"),
            ("ORDER_VALIDATED", "temporal", "Workflow đã pickup task"),
            ("INVENTORY_RESERVED", "inventory-service", "gRPC internal call OK"),
            ("PAYMENT_APPROVED", "temporal", "Mock payment activity OK"),
            ("ORDER_COMPLETED", "temporal", "Workflow kết thúc happy path"),
        ],
        "#ffe3df",
    )
    live_result = ASSETS / "live-smoke-result.json"
    if live_result.exists():
        try:
            data = json.loads(live_result.read_text(encoding="utf-8-sig"))
            happy = data.get("happy_order", {})
            failure = data.get("failure_order", {})
            happy_summary = happy.get("audit_summary", {})
            failure_summary = failure.get("audit_summary", {})
            make_terminal_panel(
                ASSETS / "live-api-smoke-result.png",
                "Live smoke test result",
                [
                    "GET /healthz -> { status: ok }",
                    f"Happy order: {happy.get('id', '')}",
                    f"  status: {happy.get('status', '')}",
                    f"  workflow: {happy.get('workflow', {}).get('status', '')}",
                    f"  audit events: {happy_summary.get('count', '')}, last: {happy_summary.get('last_type', '')}",
                    "",
                    f"Failure order: {failure.get('id', '')}",
                    f"  status: {failure.get('status', '')}",
                    f"  workflow: {failure.get('workflow', {}).get('status', '')}",
                    f"  audit events: {failure_summary.get('count', '')}, last: {failure_summary.get('last_type', '')}",
                    "",
                    "Signals: API response, DB state, Temporal history, audit timeline.",
                ],
            )
        except Exception:
            pass


def setup_doc(title):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(20, 40, 52)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(20, 40, 52)
    doc.add_paragraph("Bản nháp Note2Node - văn phong intern, xưng mình").alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def add_para(doc, text):
    for part in text.strip().split("\n\n"):
        doc.add_paragraph(part.strip())


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def add_image(doc, filename, caption, width=6.4):
    doc.add_picture(str(ASSETS / filename), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)


def pick_asset(primary, fallback):
    return primary if (ASSETS / primary).exists() else fallback


def add_mermaid_source(doc, filename):
    doc.add_paragraph("Mermaid source để reviewer có thể copy/chỉnh lại nếu cần:")
    add_code(doc, (ASSETS / filename).read_text(encoding="utf-8"))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def article_1():
    doc = setup_doc("Từ một email tech stack đến playground chạy thật: học Go Gin, Keycloak, Temporal, MySQL, MongoDB và gRPC bằng một mini Order System")
    doc.add_heading("1. Vì sao mình dựng playground này", level=1)
    add_para(doc, """
    Khi nhận một danh sách tech stack gồm Go Gin, MySQL, MongoDB, Keycloak, Temporal và gRPC, phản xạ đầu tiên của mình là mở từng docs ra đọc. Nhưng đọc từng món riêng lẻ khá nhanh bị rơi vào cảm giác: mình biết định nghĩa, nhưng không hình dung được hệ thống chạy thật sẽ đi qua những bước nào.

    Vì vậy mình chọn dựng một playground nhỏ: một mini Order Processing System. Mục tiêu của bài này không phải làm một hệ thống production, mà là có một flow đủ thật để mình nhìn thấy auth, database, workflow và service-to-service call phối hợp với nhau.
    """)
    doc.add_heading("2. Cần chuẩn bị gì trước khi chạy", level=1)
    add_bullets(doc, [
        "Docker Desktop và Docker Compose để chạy local environment.",
        "PowerShell để gọi API và chạy lệnh kiểm tra.",
        "Go cơ bản: package, module, HTTP handler.",
        "HTTP API/JWT cơ bản: biết Bearer token được gửi qua header Authorization.",
    ])
    doc.add_heading("3. Đọc nhanh trước khi vào flow", level=1)
    add_para(doc, """
    Bài này có nhiều tên công nghệ, nhưng mình sẽ đọc chúng như vai trò trong một ca trực nhỏ. Gin là cửa nhận request. Keycloak là nơi hỏi "user này là ai và có role gì?". MySQL là nơi giữ trạng thái order chính. MongoDB là cuốn sổ audit để biết order đã đi qua bước nào. Temporal là người điều phối các bước sau khi API trả response. gRPC là đường gọi nội bộ giữa các service.

    Nói ngắn gọn: đừng học từng chữ trước. Hãy nhìn một order đi từ PENDING sang COMPLETED, rồi gắn từng công nghệ vào đúng chỗ nó xuất hiện.
    """)
    doc.add_heading("4. Mình sẽ build gì", level=1)
    add_para(doc, """
    Playground có một gateway API viết bằng Go Gin. User login qua Keycloak để lấy access token, gọi API tạo order, gateway ghi dữ liệu vào MySQL và MongoDB, rồi start Temporal workflow. Workflow tiếp tục gọi inventory-service và notifier-service qua gRPC.

    Nếu mọi thứ thành công, order đi từ PENDING sang COMPLETED. Nếu inventory reject, workflow ghi lại failure và order chuyển sang FAILED.
    """)
    add_image(doc, "architecture.png", "Hình 1. Architecture tổng quan của playground.")
    add_mermaid_source(doc, "architecture.mmd")
    doc.add_heading("5. Các lựa chọn và trade-off", level=1)
    add_table(doc, ["Decision", "Vì sao mình chọn", "Trade-off"], [
        ("Order System thay vì hello world", "Có auth, DB, audit, workflow và gRPC trong cùng một flow.", "Nặng hơn demo đơn giản, nhưng học được nhiều mối nối hơn."),
        ("MySQL cho business data", "Order và order_items có cấu trúc rõ, hợp dữ liệu quan hệ.", "Cần để ý connection pool và schema migration."),
        ("MongoDB cho audit/log", "Audit event linh hoạt metadata, dễ lưu dạng document.", "Nếu query phức tạp quá thì cần thiết kế index cẩn thận."),
        ("HTTP/JSON public, gRPC internal", "Client dễ gọi API JSON, service nội bộ có contract rõ hơn bằng gRPC.", "gRPC cần tooling riêng để debug tiện."),
        ("Temporal cho workflow", "Workflow nhiều bước, có retry và history để debug.", "Workflow code cần giữ tính deterministic và versioning cẩn thận."),
    ])
    doc.add_heading("6. Build và test từng bước", level=1)
    doc.add_heading("Bước 1 - Start stack", level=2)
    add_code(doc, r"""
cd D:\playgrounds\order-playground
docker compose -f infra\docker\docker-compose.yml --env-file .env up -d --build
""")
    add_para(doc, "Sau khi start, những màn hình mình hay mở nhất là Gateway ở port 8080, Keycloak ở 8081 và Temporal UI ở 8088. Trong lần test thật của mình, máy đang có MongoDB và Temporal khác giữ port 27017/7233, nên mình chạy Mongo ở 27018 và expose Temporal service ra 7234 bằng compose override. Đây là lỗi local rất đời thường: app chưa chắc sai, đôi khi chỉ là hai container cùng muốn dùng một port.")
    add_image(doc, "docker-containers.png", "Hình 2. Các container/port chính trong playground.", 6.3)
    add_image(doc, pick_asset("live-api-smoke-result.png", "smoke-test-summary.png"), "Hình 3. Smoke test chạy thật: một order COMPLETED và một order FAILED để đối chiếu flow.")
    doc.add_heading("Bước 2 - Login lấy token", level=2)
    add_code(doc, r"""
$userToken = (Invoke-RestMethod -Method Post `
  -Uri "http://localhost:8081/realms/order-playground/protocol/openid-connect/token" `
  -ContentType "application/x-www-form-urlencoded" `
  -Body @{
    client_id="gateway-api"
    grant_type="password"
    username="user1"
    password="user1pass"
  }).access_token
""")
    add_para(doc, "Ở bước này mình chỉ cần nhớ: Keycloak là nơi phát thẻ vào cửa. Gateway không tự login user, nó chỉ nhìn thẻ đó có đúng issuer, đúng audience và đúng role hay không.")
    add_image(doc, pick_asset("live-keycloak-gateway-client.png", "keycloak-realm-screen.png"), "Hình 4. Ảnh chụp thật Keycloak: client gateway-api trong realm order-playground.", 6.3)
    doc.add_heading("Bước 3 - Tạo order", level=2)
    add_code(doc, r"""
$body = @{
  customer_name = "Nguyen Trung"
  currency = "USD"
  items = @(
    @{ sku = "BOOK-001"; quantity = 1; price = 15.5 },
    @{ sku = "PEN-002"; quantity = 2; price = 4.25 }
  )
} | ConvertTo-Json -Depth 5

$order = Invoke-RestMethod -Method Post `
  -Uri "http://localhost:8080/api/orders" `
  -Headers @{ Authorization = "Bearer $userToken" } `
  -ContentType "application/json" `
  -Body $body
""")
    add_image(doc, "create-order-sequence.png", "Hình 5. Sequence tạo order từ login đến workflow.")
    add_mermaid_source(doc, "create-order-sequence.mmd")
    doc.add_heading("Bước 4 - Kiểm tra dữ liệu", level=2)
    add_code(doc, r"""
docker exec order-playground-mysql mysql -uorder_app -porder_pass -D order_playground -e "SELECT * FROM orders;"
docker exec order-playground-mongodb mongosh --username mongoadmin --password mongopass --authenticationDatabase admin order_playground --quiet --eval "db.audit_events.find().pretty()"
""")
    add_para(doc, "MySQL giữ trạng thái nghiệp vụ chính của order. MongoDB giữ audit trail để mình đọc được order đã đi qua những mốc nào. Temporal UI thì giống camera hành trình của workflow: mở lên là thấy workflow đang chạy, đã xong, hay fail ở activity nào.")
    add_image(doc, pick_asset("live-temporal-completed-detail.png", "temporal-completed-screen.png"), "Hình 6. Ảnh chụp thật Temporal UI: workflow order đã COMPLETED.", 6.3)
    add_image(doc, "db-audit-events-screen.png", "Hình 7. Audit events dùng để đối chiếu flow.", 6.3)
    doc.add_heading("7. Góc đọc repo để không bị ngợp", level=1)
    add_para(doc, """
    Khi mới mở repo, mình không đọc từ trên xuống dưới như đọc sách, vì cách đó rất dễ bị ngợp. Mình đi theo đúng lifecycle của một request. Đầu tiên mình mở router của gateway để xem endpoint nào nhận request. Sau đó mình lần theo handler tạo order, repository MySQL, repository MongoDB, rồi tới đoạn start Temporal workflow. Khi workflow chạy, mình mới nhảy qua activity để xem vì sao nó gọi inventory-service và notifier-service bằng gRPC.

    Cách đọc này giúp mình nối được code với màn hình debug. Nếu API trả 401 hoặc 403 thì mình mở Keycloak trước. Nếu order đã tạo nhưng không đổi status thì mình mở Temporal UI. Nếu workflow báo completed nhưng dữ liệu không đúng thì mình kiểm tra MySQL và audit events trong MongoDB. Nói cách khác, repo này không chỉ để chạy demo, mà còn là một bản đồ nhỏ để luyện thói quen debug backend có nhiều thành phần.
    """)
    add_bullets(doc, [
        "Đọc public API trước: healthz, readyz, version để biết app sống chưa.",
        "Đọc protected API sau: orders và admin routes để hiểu middleware auth/role guard.",
        "Đọc workflow cuối cùng: vì Temporal có retry và chạy async, nên cần nhìn cùng lúc code, UI và DB state.",
    ])
    doc.add_heading("8. Checklist test nhanh", level=1)
    add_bullets(doc, [
        "Happy path: tạo order với quantity nhỏ hơn hoặc bằng 5, kỳ vọng order COMPLETED.",
        "Failure path: tạo order có quantity lớn hơn 5, inventory-service reject và workflow chuyển order sang FAILED.",
        "Auth path: gọi /api/orders không có token để thấy 401, dùng user1 gọi /api/admin/orders để thấy 403.",
    ])
    doc.add_heading("9. Kết bài", level=1)
    add_para(doc, """
    Sau bài này, mình không chỉ có một repo chạy được, mà còn có một flow để học tiếp. Khi nhìn vào một request tạo order, mình thấy được vì sao cần Keycloak, vì sao tách MySQL và MongoDB, vì sao Temporal không chỉ là background job, và vì sao gRPC hợp với service nội bộ.

    Bài học lớn nhất của mình là: với một stack rộng, playground nhỏ nhưng chạy thật giúp mình học nhanh hơn rất nhiều so với đọc từng tech riêng lẻ.
    """)
    add_references(doc)
    doc.save(OUT / "01-build-order-playground.docx")


def article_2():
    doc = setup_doc("Debug một Backend Playground: mình học gì từ Keycloak, Temporal, JWT issuer, workflow retry và database state")
    doc.add_heading("1. Tóm tắt nhanh", level=1)
    add_bullets(doc, [
        "Keycloak nên được nhìn như nơi debug identity/access, không chỉ là màn hình login.",
        "Temporal UI rất hữu ích khi order bị kẹt ở trạng thái RUNNING hoặc workflow retry mãi.",
        "JWT issuer mismatch là lỗi dễ gặp khi local host và Docker hostname khác nhau.",
        "Database state và audit events giúp mình kiểm chứng workflow đã đi tới bước nào.",
        "Debug các lỗi nhỏ trong playground làm mình hiểu kiến trúc hơn là chỉ đọc docs.",
    ])
    doc.add_heading("2. Vì sao debug mới làm mình hiểu", level=1)
    add_para(doc, """
    Lúc playground chạy được lần đầu, mình tưởng phần khó nhất đã xong. Nhưng thật ra phần giúp mình học nhiều nhất lại là lúc nó lỗi. Có lỗi đến từ môi trường như Docker daemon chưa chạy, có lỗi đến từ Keycloak như issuer mismatch, và có lỗi đến từ Temporal như activity name không khớp.

    Bài này không đi lại toàn bộ cách build. Mình tập trung vào cách debug và những mental model mình rút ra sau khi va vào các lỗi đó.
    """)
    doc.add_heading("3. Background cần biết", level=1)
    add_para(doc, """
    Với Keycloak, mình cần phân biệt authentication và authorization. Authentication trả lời user là ai, authorization trả lời user được làm gì. Trong playground, user lấy access token từ realm order-playground, gateway verify token rồi đọc role user/admin.

    Với Temporal, mình cần hiểu workflow là quy trình nhiều bước, activity là từng bước xử lý cụ thể, worker là process nhận task từ task queue. Khi activity fail, Temporal có thể retry và lưu lại history để mình debug.

    Với database, mình dùng MySQL cho dữ liệu nghiệp vụ chính và MongoDB cho audit events. Khi API hoặc workflow lỗi, DB thường là nơi giúp mình đối chiếu state thật.
    """)
    add_table(doc, ["Thuật ngữ", "Mình hiểu đơn giản là", "Nhìn ở đâu trong playground"], [
        ("Realm", "Một không gian quản lý user/client/role riêng trong Keycloak.", "Keycloak Admin Console"),
        ("Client", "Ứng dụng xin token hoặc dùng token. Ở đây là gateway-api.", "Keycloak > order-playground > Clients"),
        ("JWT issuer", "Dòng ghi token do ai phát hành. Sai issuer là gateway từ chối token.", "Token payload + gateway config"),
        ("Workflow", "Kịch bản xử lý order gồm nhiều bước.", "Temporal UI"),
        ("Activity", "Một bước cụ thể trong workflow, ví dụ reserve inventory.", "Temporal workflow history"),
        ("Audit event", "Dòng nhật ký nghiệp vụ để đọc lại order đã đi qua đâu.", "MongoDB audit_events"),
    ])
    doc.add_heading("4. Mental model khi debug", level=1)
    add_image(doc, "debug-tree.png", "Hình 1. Debug decision tree: gặp lỗi thì nên mở màn hình nào trước.")
    add_mermaid_source(doc, "debug-tree.mmd")
    add_para(doc, "Mình chia debugging thành ba đường nhìn: request path, control path và debug path. Request path đi từ user đến gateway rồi DB/workflow. Control path là Temporal worker xử lý activity. Debug path là mở đúng công cụ: Keycloak cho auth, Temporal cho workflow, DB cho state, logs cho startup/runtime.")
    add_para(doc, "Good debugging trong playground này là mỗi lần chỉ kiểm chứng một giả thuyết. Ví dụ: nếu API trả 403, mình không sửa workflow ngay. Mình kiểm tra token có role gì trước. Nếu order đứng ở PENDING, mình không sửa Keycloak nữa. Mình mở Temporal history để xem activity nào đang retry.")
    add_bullets(doc, [
        "Bad: đổi config Keycloak, restart Temporal, sửa code repository cùng lúc rồi hy vọng hết lỗi.",
        "Good: ghi lại symptom, chọn một màn hình để kiểm chứng, sau đó mới sửa đúng chỗ.",
    ])
    doc.add_heading("5. Những lỗi thật mình gặp", level=1)
    add_image(doc, "debug-incidents.png", "Hình 2. Các lỗi thật đã gặp trong quá trình dựng playground.")
    doc.add_heading("PowerShell chặn .ps1", level=2)
    add_para(doc, "Lỗi này không liên quan app. Windows chặn script do execution policy. Cách vòng qua nhanh là dùng ExecutionPolicy Bypass hoặc chạy lệnh Invoke-RestMethod inline. Bài học của mình: lỗi demo đôi khi nằm ở môi trường test trước khi nằm ở code.")
    doc.add_heading("Docker daemon chưa chạy", level=2)
    add_para(doc, "Khi docker ps báo không tìm thấy Docker engine pipe, mình biết container không phải nguyên nhân. Cần kiểm tra Docker Desktop trước, rồi mới debug service.")
    doc.add_heading("MySQL Public Key Retrieval is not allowed", level=2)
    add_para(doc, "Lỗi này thường đến từ GUI/JDBC client kết nối MySQL 8. Fix nhanh là thêm allowPublicKeyRetrieval=true&useSSL=false vào JDBC URL. Container MySQL không hỏng; client chỉ chưa được phép lấy public key.")
    doc.add_heading("Keycloak master vs order-playground", level=2)
    add_image(doc, "keycloak-model.png", "Hình 3. Keycloak mental model: master realm khác order-playground realm.")
    add_mermaid_source(doc, "keycloak-model.mmd")
    add_image(doc, pick_asset("live-keycloak-master-realm-client.png", "keycloak-clients-screen.png"), "Hình 4. Ảnh chụp thật: client order-playground-realm nằm trong master realm, dễ gây nhầm với business client.", 6.3)
    add_para(doc, "Trong Keycloak, master là realm quản trị. order-playground là realm của app. Client gateway-api mới là client business thật. Client order-playground-realm trong master là client quản trị nội bộ, không phải client để app lấy token.")
    doc.add_heading("JWT issuer mismatch", level=2)
    add_para(doc, "Token lấy từ browser/PowerShell có issuer dạng localhost:8081, trong khi gateway trong Docker có thể fetch Keycloak qua hostname keycloak:8080. Nếu gateway dùng luôn hostname nội bộ làm issuer, token sẽ bị invalid issuer. Fix là tách KEYCLOAK_ISSUER_URL và KEYCLOAK_JWKS_BASE_URL.")
    doc.add_heading("Temporal namespace và activity mismatch", level=2)
    add_image(doc, pick_asset("live-temporal-ui.png", "temporal-flow.png"), "Hình 5. Ảnh chụp thật Temporal UI: cùng lúc nhìn được Running, Completed và Failed workflows.", 6.3)
    add_image(doc, "temporal-flow.png", "Hình 6. Temporal flow và failure path.")
    add_mermaid_source(doc, "temporal-retry.mmd")
    add_para(doc, """
    Temporal auto-setup cần thời gian để tạo namespace default. Nếu gateway start worker quá sớm, worker fail. Mình thêm retry startup để gateway kiên nhẫn đợi Temporal.

    Một lỗi khác thú vị hơn là workflow schedule activity tên Activities.ValidateOrderActivity, nhưng worker chỉ register ValidateOrderActivity. Temporal UI cho thấy Pending Activity và retry. Sau khi đổi tên gọi activity cho khớp, workflow chạy từ PENDING sang COMPLETED.
    """)
    doc.add_heading("6. Trade-off mình rút ra", level=1)
    add_table(doc, ["Chủ đề", "Điểm tốt", "Trade-off cần nhớ"], [
        ("Stateless JWT", "Verify nhanh, microservice không cần gọi auth server mỗi request.", "Khó revoke tức thì nếu token còn hạn, cần blacklist/introspection cho API nhạy cảm."),
        ("Temporal", "Có retry, state bền vững, workflow history rõ.", "Workflow code cần versioning, tránh thay đổi non-deterministic."),
        ("gRPC internal", "Contract rõ, hợp service-to-service.", "Debug thủ công khó hơn REST nếu thiếu tooling."),
        ("Docker Compose local", "Dễ dựng lab, dễ học dependency giữa service.", "Không đại diện đầy đủ production như Kubernetes/managed services."),
    ])
    doc.add_heading("7. Bảng troubleshooting", level=1)
    add_table(doc, ["Lỗi", "Mở ở đâu", "Kiểm tra gì", "Bài học"], [
        ("401/403", "Keycloak + gateway logs", "realm, client, token, role", "Auth là cả identity lẫn permission."),
        ("Order PENDING lâu", "Temporal UI", "workflow history, pending activity, retry", "Workflow state không nằm trong HTTP request."),
        ("Không thấy audit", "MongoDB", "audit_events theo orderId", "Audit trail giúp đọc lại flow."),
        ("Không thấy order", "MySQL", "orders, order_items, workflow_runs", "Business state cần một source of truth rõ."),
        ("Container không lên", "Docker logs", "port conflict, healthcheck, startup retry", "Debug hạ tầng trước khi sửa code."),
    ])
    doc.add_heading("8. Tổng kết và bước tiếp theo", level=1)
    add_para(doc, """
    Sau khi debug playground này, mình thấy Keycloak và Temporal không còn là hai màn hình xa lạ nữa. Keycloak giúp mình trả lời vì sao user không vào được. Temporal giúp mình trả lời vì sao workflow chưa xong. MySQL và MongoDB giúp mình kiểm chứng state thật thay vì chỉ tin response API.

    Hướng tiếp theo mình muốn thử là tách Temporal worker ra service riêng, thêm Postman collection, thử workflow versioning và sau đó mới mở rộng sang RabbitMQ, Redis hoặc Kubernetes. Làm từng bước nhỏ như vậy giúp mình không bị biến bài học thành một danh sách tech stack quá rộng.
    """)
    add_references(doc)
    doc.save(OUT / "02-debug-learn-keycloak-temporal.docx")


def add_references(doc):
    doc.add_heading("References", level=1)
    refs = [
        "Gin Quickstart: https://gin-gonic.com/en/docs/quickstart/",
        "Go database/sql - Managing connections: https://go.dev/doc/database/manage-connections",
        "MongoDB Go Driver Docs: https://www.mongodb.com/docs/drivers/go/v2.0/",
        "Temporal Docs: https://docs.temporal.io/",
        "Temporal Go SDK package docs: https://pkg.go.dev/go.temporal.io/sdk/workflow",
        "Keycloak Server Administration Guide: https://www.keycloak.org/docs/latest/server_admin/",
        "OpenID Connect Core 1.0: https://openid.net/specs/openid-connect-core-1_0-18.html",
        "OAuth 2.0 RFC 6749: https://datatracker.ietf.org/doc/rfc6749/",
        "gRPC Introduction: https://grpc.io/docs/what-is-grpc/introduction/",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    make_assets()
    article_1()
    article_2()
    print(f"Generated: {OUT / '01-build-order-playground.docx'}")
    print(f"Generated: {OUT / '02-debug-learn-keycloak-temporal.docx'}")
