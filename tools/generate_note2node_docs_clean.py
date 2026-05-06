# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\playgrounds\order-playground")
OUT = ROOT / "docs" / "note2node"
ASSETS = OUT / "assets"


ARTICLE_1 = OUT / "01-build-order-playground-outline-clean.docx"
ARTICLE_2 = OUT / "02-debug-learn-keycloak-temporal-outline-clean.docx"


def asset(primary, fallback):
    return primary if (ASSETS / primary).exists() else fallback


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def setup_doc(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(3)

    for name, size, before, after in [
        ("Heading 1", 13, 10, 3),
        ("Heading 2", 11.2, 7, 2),
        ("Heading 3", 10.5, 5, 1),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.bold = True
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(18, 38, 50)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(18, 38, 50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(80, 91, 103)
    return doc


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.6)
    return p


def image(doc, filename, caption, width=5.75):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(ASSETS / filename), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(90, 99, 110)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "EAF4F0")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.7)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(8.5)
    if widths:
        for row in tbl.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return tbl


def references(doc):
    doc.add_heading("References", level=1)
    bullets(doc, [
        "Gin Quickstart: https://gin-gonic.com/en/docs/quickstart/",
        "Go database/sql - Managing connections: https://go.dev/doc/database/manage-connections",
        "MongoDB Go Driver Docs: https://www.mongodb.com/docs/drivers/go/current/",
        "Temporal Docs: https://docs.temporal.io/",
        "Keycloak Server Administration Guide: https://www.keycloak.org/docs/latest/server_admin/",
        "OpenID Connect Core 1.0: https://openid.net/specs/openid-connect-core-1_0.html",
        "OAuth 2.0 RFC 6749: https://datatracker.ietf.org/doc/rfc6749/",
        "gRPC Introduction: https://grpc.io/docs/what-is-grpc/introduction/",
    ])


def article_1():
    doc = setup_doc(
        "Từ một email tech stack đến playground chạy thật: học Go Gin, Keycloak, Temporal, MySQL, MongoDB và gRPC bằng một mini Order System",
        "Bản nháp Note2Node - tutorial backend-only, build -> debug -> learn",
    )

    doc.add_heading("1. Mở bài: vì sao mình không học từng món rời rạc", level=1)
    para(doc, "Mình bắt đầu từ một email liệt kê khá nhiều thứ: Go Gin, MySQL, MongoDB, Temporal, Keycloak, OIDC/OAuth, HTTP/1.1, HTTP/2 và gRPC. Nếu mở từng docs ra đọc riêng, mình có thể nhớ định nghĩa, nhưng rất khó hình dung khi chúng đứng chung trong một backend thật thì request sẽ đi qua đâu.")
    para(doc, "Vì vậy mình chọn dựng một playground nhỏ: một mini Order System. Mục tiêu không phải làm production-ready system, mà là có một flow chạy được để vừa build, vừa debug, vừa tự trả lời: công nghệ này xuất hiện ở bước nào và giải quyết vấn đề gì?")

    doc.add_heading("2. Prerequisites: cần biết gì trước khi chạy", level=1)
    bullets(doc, [
        "Docker Desktop và Docker Compose để chạy local environment.",
        "PowerShell cơ bản để gọi API, lấy token và query nhanh.",
        "Go cơ bản: package, handler, context và cách app đọc env config.",
        "HTTP/JWT cơ bản: biết Bearer token được gửi qua header Authorization.",
        "Không cần biết sâu Temporal hay Keycloak trước. Bài này sẽ dùng flow order để giải thích vừa đủ.",
    ])

    doc.add_heading("3. What we build: mini Order Processing System", level=1)
    para(doc, "Flow chính rất nhỏ: user login qua Keycloak, gọi API tạo order vào Gin gateway, gateway ghi order vào MySQL, ghi audit event vào MongoDB, rồi start Temporal workflow. Workflow sau đó gọi inventory-service và notifier-service qua gRPC, cuối cùng cập nhật trạng thái order thành COMPLETED hoặc FAILED.")
    para(doc, "Cách mình tự nhớ vai trò từng thành phần là: Gin nhận request, Keycloak cấp token, MySQL giữ state chính, MongoDB giữ nhật ký, Temporal điều phối việc dài hơi, còn gRPC là đường nói chuyện nội bộ giữa service.")

    doc.add_heading("4. Architecture overview", level=1)
    image(doc, "architecture.png", "Hình 1. Architecture tổng quan của playground.", 5.85)
    para(doc, "Public API dùng HTTP/JSON vì dễ gọi bằng curl/Postman. Internal service dùng gRPC vì contract rõ hơn và chạy trên HTTP/2. Temporal không nằm trong request trực tiếp; nó nhận workflow sau khi gateway đã tạo order PENDING.")
    para(doc, "Mình cũng cố tình chia ranh giới giao thức cho dễ học. Từ user vào gateway là HTTP/1.1 JSON, vì đây là kiểu API quen thuộc nhất để test. Từ workflow sang inventory/notifier là gRPC, để thấy HTTP/2 xuất hiện trong giao tiếp nội bộ. Còn Keycloak dùng OIDC/OAuth2 để phát token, không phải nơi lưu order hay chạy workflow.")
    para(doc, "Mermaid source của architecture và sequence diagram được lưu trong thư mục assets để reviewer có thể copy/chỉnh lại nếu cần.")

    doc.add_heading("5. Decision & trade-off", level=1)
    table(doc, ["Decision", "Vì sao chọn", "Trade-off"], [
        ("Order System thay vì hello world", "Một flow nhỏ nhưng có auth, DB, audit, workflow và internal service call.", "Setup nặng hơn, nhưng bù lại thấy được mối nối thật giữa các tech."),
        ("MySQL cho order", "Order và order_items có schema rõ, hợp relational database.", "Phải để ý migration và connection pool."),
        ("MongoDB cho audit", "Audit event linh hoạt metadata, dễ lưu dạng document.", "Muốn query tốt vẫn phải nghĩ tới index."),
        ("HTTP/JSON public, gRPC internal", "Client bên ngoài dễ test, service bên trong có contract chặt hơn.", "gRPC debug khó hơn REST nếu thiếu tooling."),
        ("Temporal cho workflow", "Có retry, history và UI để quan sát order đi qua từng bước.", "Workflow code cần cẩn thận với versioning/determinism."),
    ], widths=[1.45, 2.45, 2.15])

    doc.add_heading("6. Step-by-step: chạy playground", level=1)
    doc.add_heading("6.1 Start Docker Compose", level=2)
    code(doc, r"""
cd D:\playgrounds\order-playground
docker compose -f infra\docker\docker-compose.yml --env-file .env up -d --build
""")
    para(doc, "Trong lần test thật trên máy mình, port 27017 và 7233 đang bị container khác giữ, nên mình dùng override để Mongo expose ra 27018 và Temporal expose ra 7234. Đây là bài học local khá thực tế: trước khi nghi code sai, hãy kiểm tra port conflict.")
    image(doc, "docker-containers.png", "Hình 2. Các container/port chính trong playground.", 5.65)
    image(doc, asset("live-api-smoke-result.png", "smoke-test-summary.png"), "Hình 3. Smoke test thật: một order COMPLETED và một order FAILED.", 5.65)

    doc.add_heading("6.2 Login lấy token từ Keycloak", level=2)
    code(doc, r"""
$userToken = (Invoke-RestMethod -Method Post `
  -Uri "http://localhost:8081/realms/order-playground/protocol/openid-connect/token" `
  -ContentType "application/x-www-form-urlencoded" `
  -Body @{
    client_id="gateway-api"
    grant_type="password"
    username="user1"
    password="user1pass"
  }).access_token
""")
    para(doc, "Ở đây Keycloak giống nơi phát thẻ vào cửa. Gateway không tự login user, nó chỉ kiểm tra thẻ đó có đúng issuer, đúng audience và có role phù hợp không.")
    image(doc, asset("live-keycloak-gateway-client.png", "keycloak-realm-screen.png"), "Hình 4. Keycloak: client gateway-api trong realm order-playground.", 3.75)

    doc.add_heading("6.3 Gọi POST /api/orders", level=2)
    code(doc, r"""
$body = @{
  customer_name = "Nguyen Trung"
  currency = "USD"
  items = @(
    @{ sku = "BOOK-001"; quantity = 1; price = 15.5 },
    @{ sku = "PEN-002"; quantity = 2; price = 4.25 }
  )
} | ConvertTo-Json -Depth 5

$order = Invoke-RestMethod -Method Post `
  -Uri "http://localhost:8080/api/orders" `
  -Headers @{ Authorization = "Bearer $userToken" } `
  -ContentType "application/json" `
  -Body $body
""")
    image(doc, "create-order-sequence.png", "Hình 5. Sequence tạo order từ login đến workflow.", 5.75)

    doc.add_heading("6.4 Kiểm tra MySQL, MongoDB và Temporal", level=2)
    code(doc, r"""
docker exec order-playground-mysql mysql -uorder_app -porder_pass -D order_playground -e "SELECT id,status,created_by FROM orders;"
docker exec order-playground-mongodb mongosh --username mongoadmin --password mongopass --authenticationDatabase admin order_playground --quiet --eval "db.audit_events.find().pretty()"
""")
    para(doc, "MySQL cho mình biết state hiện tại của order. MongoDB cho mình timeline nghiệp vụ. Temporal UI cho mình biết workflow fail ở bước nào, retry bao nhiêu lần, hay đã completed.")
    para(doc, "Một mẹo nhỏ khi test là đừng chỉ nhìn response của POST /api/orders. Response lúc đầu trả PENDING là đúng, vì workflow chạy async. Mình cần chờ vài giây rồi gọi GET /api/orders/:id hoặc mở Temporal UI để biết kết quả cuối cùng. Đây cũng là chỗ mình hiểu rõ hơn sự khác nhau giữa request lifecycle của Gin và workflow lifecycle của Temporal.")
    image(doc, asset("live-temporal-completed-detail.png", "temporal-completed-screen.png"), "Hình 6. Temporal UI: workflow order đã COMPLETED.", 5.25)
    image(doc, "db-audit-events-screen.png", "Hình 7. Audit events để đối chiếu flow.", 5.65)

    doc.add_heading("7. Testing: happy path và failure path", level=1)
    bullets(doc, [
        "Happy path: tạo order với quantity nhỏ hơn hoặc bằng 5, kỳ vọng order đi từ PENDING sang COMPLETED.",
        "Failure path: tạo order có quantity lớn hơn 5, inventory-service reject và workflow chuyển order sang FAILED.",
        "Auth path: gọi /api/orders không token để thấy 401, dùng user1 gọi /api/admin/orders để thấy 403.",
    ])

    doc.add_heading("8. Cách đọc repo để không bị ngợp", level=1)
    para(doc, "Nếu đọc repo từ đầu đến cuối, mình rất dễ bị lạc vì có nhiều folder. Cách hợp lý hơn là đọc theo hành trình của một order. Mỗi lần đi qua một lớp, mình chỉ hỏi một câu: lớp này nhận gì, ghi gì, gọi ai tiếp theo?")
    numbered(doc, [
        "Đọc router/handler trước để biết request vào đâu.",
        "Lần theo repository MySQL/MongoDB để biết dữ liệu được ghi ở đâu.",
        "Mở workflow/activity sau cùng để hiểu phần async.",
        "Khi lỗi auth thì mở Keycloak; khi order kẹt thì mở Temporal; khi state sai thì mở DB.",
    ])

    doc.add_heading("9. Wrap-up", level=1)
    para(doc, "Điểm mình thích ở playground này là nó không cố dạy tất cả lý thuyết cùng lúc. Nó cho mình một order thật để lần theo. Từ đó, mỗi thuật ngữ không còn đứng riêng: Keycloak nằm ở cửa vào, MySQL/MongoDB nằm ở state/audit, Temporal nằm ở workflow, còn gRPC nằm ở service nội bộ.")
    para(doc, "Bài học lớn nhất của mình: với một stack rộng, một playground nhỏ nhưng chạy thật giúp học nhanh hơn nhiều so với đọc từng tech riêng lẻ.")
    references(doc)
    doc.save(ARTICLE_1)


def article_2():
    doc = setup_doc(
        "Debug một Backend Playground: mình học gì từ Keycloak, Temporal, JWT issuer, workflow retry và database state",
        "Bản nháp Note2Node - deep-dive theo incident, debug -> learn",
    )

    doc.add_heading("1. Executive Summary", level=1)
    para(doc, "Sau khi build playground, phần làm mình hiểu sâu nhất không phải lúc mọi thứ chạy xanh, mà là lúc nó lỗi. Bài này ghi lại cách mình debug các lỗi thật: PowerShell chặn script, Docker/port conflict, MySQL client lỗi public key, Keycloak realm gây nhầm, JWT issuer mismatch và Temporal activity retry.")
    bullets(doc, [
        "Keycloak nên được xem như nơi debug identity/access, không chỉ là màn hình login.",
        "Temporal UI là nơi xem workflow history, activity retry và trạng thái order async.",
        "DB state và audit events là bằng chứng để kiểm tra workflow đã đi tới đâu.",
        "Debug tốt là giảm uncertainty từng bước, không sửa nhiều thứ cùng lúc.",
    ])

    doc.add_heading("2. Motivation: hệ thống chạy được chưa chắc đã hiểu", level=1)
    para(doc, "Lúc API trả 200, mình dễ có cảm giác đã hiểu hệ thống. Nhưng chỉ cần token bị reject, order kẹt PENDING, hoặc workflow retry mãi là mình nhận ra: hiểu thật nghĩa là biết mở đúng màn hình, nhìn đúng signal, rồi sửa đúng chỗ.")
    para(doc, "Vì vậy bài này không liệt kê lại tech stack. Mình đi theo incident: symptom là gì, mình mở đâu, kiểm tra gì, và rút ra mental model nào.")

    doc.add_heading("3. Background knowledge cần biết", level=1)
    table(doc, ["Thuật ngữ", "Mình hiểu đơn giản là", "Nhìn ở đâu"], [
        ("Realm", "Không gian riêng để quản lý user, client và role.", "Keycloak Admin Console"),
        ("Client", "Ứng dụng xin token hoặc dùng token. Ở đây là gateway-api.", "Keycloak > Clients"),
        ("JWT issuer", "Dòng ghi token do ai phát hành. Sai issuer thì gateway reject.", "Token payload + gateway config"),
        ("Workflow", "Kịch bản nhiều bước xử lý order.", "Temporal UI"),
        ("Activity", "Một bước cụ thể, ví dụ reserve inventory.", "Temporal workflow history"),
        ("Audit event", "Nhật ký nghiệp vụ để đọc lại order đã đi qua đâu.", "MongoDB audit_events"),
    ], widths=[1.35, 3.0, 1.65])
    para(doc, "Với Keycloak, mình tách authentication và authorization: user là ai, và user được làm gì. Với Temporal, mình tách workflow và activity: quy trình lớn và từng bước nhỏ. Với DB, mình tách business state và audit trail: trạng thái hiện tại và lịch sử đã xảy ra.")

    doc.add_heading("4. Mental model: mở đúng màn hình trước", level=1)
    image(doc, "debug-tree.png", "Hình 1. Debug decision tree: gặp lỗi thì mở màn hình nào trước.", 5.65)
    para(doc, "Good debugging trong playground này là mỗi vòng chỉ kiểm chứng một giả thuyết. Nếu API trả 403, mình kiểm tra token/role trước. Nếu order đứng PENDING, mình mở Temporal history trước. Nếu workflow completed nhưng dữ liệu sai, mình kiểm tra MySQL và MongoDB.")
    table(doc, ["Path", "Nó trả lời câu hỏi gì?", "Tool mình mở"], [
        ("Request path", "Request từ user đi qua gateway, auth, DB và start workflow như thế nào?", "API response + gateway logs"),
        ("Control path", "Temporal worker đã nhận task chưa, activity nào đang retry?", "Temporal UI"),
        ("Debug path", "State thật đang nằm ở đâu và có khớp với kỳ vọng không?", "MySQL, MongoDB, Docker logs"),
    ], widths=[1.35, 3.2, 1.45])
    bullets(doc, [
        "Bad: đổi Keycloak config, restart Temporal và sửa repository cùng lúc.",
        "Good: ghi symptom, chọn một signal, test lại, rồi mới sửa đúng chỗ.",
    ])

    doc.add_heading("5. Deep dive by incident", level=1)
    image(doc, "debug-incidents.png", "Hình 2. Các lỗi thật trong quá trình dựng playground.", 4.95)
    para(doc, "Mỗi incident bên dưới mình đọc theo cùng một format: symptom xuất hiện ở đâu, mình mở tool nào trước, signal nào xác nhận giả thuyết, và bài học rút ra là gì. Format này giúp bài không biến thành danh sách lỗi rời rạc, mà vẫn giữ được mạch debug -> learn.")

    doc.add_heading("5.1 PowerShell chặn .ps1", level=2)
    para(doc, "Symptom: chạy script lấy token bị báo running scripts is disabled. Đây không phải lỗi Keycloak. Đây là policy của Windows. Cách xử lý nhanh là chạy PowerShell với ExecutionPolicy Bypass hoặc dùng Invoke-RestMethod inline.")

    doc.add_heading("5.2 Docker daemon và port conflict", level=2)
    para(doc, "Symptom: container không start hoặc báo port đã được allocate. Lần test thật của mình có MongoDB khác giữ 27017 và Temporal khác giữ 7233. Bài học: trước khi sửa code, kiểm tra docker ps và port mapping.")

    doc.add_heading("5.3 MySQL Public Key Retrieval is not allowed", level=2)
    para(doc, "Symptom: GUI/JDBC client báo Public Key Retrieval is not allowed. Container MySQL không chết. Lỗi nằm ở cách client kết nối MySQL 8. Fix nhanh là thêm allowPublicKeyRetrieval=true&useSSL=false vào JDBC URL khi dùng client dev.")

    doc.add_heading("5.4 Keycloak master vs order-playground", level=2)
    image(doc, "keycloak-model.png", "Hình 3. Mental model: master realm khác realm của app.", 5.65)
    image(doc, asset("live-keycloak-master-realm-client.png", "keycloak-clients-screen.png"), "Hình 4. Keycloak master có client order-playground-realm, không phải business client của gateway.", 4.75)
    para(doc, "Chỗ này dễ nhầm. master là realm quản trị Keycloak. order-playground là realm của app demo. Client gateway-api trong realm order-playground mới là client business dùng để lấy token. Client order-playground-realm trong master là client quản trị nội bộ do Keycloak tạo.")

    doc.add_heading("5.5 JWT issuer mismatch", level=2)
    para(doc, "Symptom: token lấy từ localhost nhưng gateway trong Docker lại dùng hostname keycloak để fetch JWKS. Nếu dùng sai issuer, token hợp lệ vẫn bị reject. Fix là tách KEYCLOAK_ISSUER_URL cho issuer public và KEYCLOAK_JWKS_BASE_URL cho đường gọi nội bộ trong Docker.")
    para(doc, "Đây là lỗi làm mình nhớ lâu nhất vì nhìn qua tưởng token sai. Thực ra token đúng, nhưng người kiểm tra token đang kỳ vọng một issuer khác. Trong local Docker, cùng một Keycloak có thể được nhìn bằng hai cái tên: localhost từ máy mình, và keycloak từ container.")

    doc.add_heading("5.6 Temporal namespace và activity mismatch", level=2)
    para(doc, "Temporal auto-setup cần thời gian để namespace default sẵn sàng. Nếu worker start quá sớm, gateway nên retry thay vì chết ngay. Một lỗi khác là activity name mismatch: workflow gọi tên activity không khớp với tên worker register, làm activity pending/retry mãi. Temporal UI giúp mình thấy điều này rõ hơn log thường.")
    para(doc, "Điểm hay của Temporal là lỗi không biến mất sau một dòng log. Nó nằm trong workflow history. Mình có thể mở lại từng event để biết activity nào được scheduled, activity nào started, activity nào failed. Với một người mới học như mình, UI này làm workflow bớt trừu tượng hơn nhiều.")
    image(doc, asset("live-temporal-ui.png", "temporal-flow.png"), "Hình 5. Temporal UI: nhìn được Running, Completed và Failed workflows.", 5.1)
    image(doc, "temporal-flow.png", "Hình 6. Failure path của Temporal workflow.", 5.1)

    doc.add_heading("6. Trade-offs", level=1)
    table(doc, ["Chủ đề", "Điểm tốt", "Trade-off"], [
        ("Stateless JWT", "Verify nhanh, gateway không phải gọi auth server mỗi request.", "Revoke không tức thì nếu token còn hạn."),
        ("Temporal", "Có retry, history và trạng thái workflow rõ.", "Workflow code cần versioning cẩn thận."),
        ("gRPC internal", "Contract rõ, hợp service-to-service.", "Debug thủ công khó hơn REST nếu thiếu tooling."),
        ("Docker Compose local", "Dễ dựng lab và học dependency giữa service.", "Không phản ánh đầy đủ production."),
    ], widths=[1.45, 2.3, 2.1])

    doc.add_heading("7. Troubleshooting table", level=1)
    table(doc, ["Lỗi", "Mở ở đâu", "Kiểm tra gì", "Bài học"], [
        ("401/403", "Keycloak + gateway logs", "realm, client, token, role", "Auth là identity + permission."),
        ("Order PENDING lâu", "Temporal UI", "workflow history, pending activity, retry", "Workflow state không nằm trong HTTP request."),
        ("Không thấy audit", "MongoDB", "audit_events theo orderId", "Audit trail giúp đọc lại flow."),
        ("Không thấy order", "MySQL", "orders, order_items, workflow_runs", "Business state cần source of truth rõ."),
        ("Container không lên", "Docker logs", "port conflict, healthcheck, startup retry", "Debug hạ tầng trước khi sửa code."),
    ], widths=[1.3, 1.45, 1.75, 1.8])

    doc.add_heading("8. Summary & next steps", level=1)
    para(doc, "Sau khi debug playground này, mình thấy Keycloak và Temporal không chỉ là hai màn hình xa lạ nữa. Keycloak giúp mình trả lời ai vào hệ thống và có quyền gì. Temporal giúp mình trả lời order đang kẹt ở bước nào. MySQL và MongoDB giúp mình kiểm chứng state thật thay vì chỉ tin response API.")
    para(doc, "Hướng tiếp theo mình muốn thử là tách Temporal worker ra service riêng, thêm Postman collection, thử workflow versioning và sau đó mở rộng sang RabbitMQ, Redis hoặc Kubernetes. Nhưng bước đầu tiên vẫn là giữ playground nhỏ, chạy thật, và debug được từng mối nối.")
    references(doc)
    doc.save(ARTICLE_2)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    article_1()
    article_2()
    print(f"Generated: {ARTICLE_1}")
    print(f"Generated: {ARTICLE_2}")


if __name__ == "__main__":
    main()
